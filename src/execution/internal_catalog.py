"""Local company catalog lookup: booking requirements vs Excel/Word/PDF/CSV (+ optional signed-in Sheet)."""

from __future__ import annotations

import csv
import json
import os
import re
from datetime import datetime, timedelta
from io import StringIO
from typing import Any, Dict, Iterable, List, Optional, Tuple

from src.execution.local_auth import load_merged_config
from src.ingestion.secret_redactor import redact_secrets


DEFAULT_REL = os.path.join("storage", "internal_data")
MAX_FILES = 40
MAX_FILE_BYTES = 2_000_000
MAX_ROWS = 5000
MAX_HITS = 8
SUPPORTED = {".csv", ".tsv", ".json", ".txt", ".md", ".xlsx", ".pdf", ".docx"}

STAY_TYPES = (
    "hotel",
    "villa",
    "apartment",
    "hostel",
    "resort",
    "cottage",
    "bnb",
    "b&b",
    "guesthouse",
    "room",
    "accommodation",
)

COUNTRY_HINTS = (
    "italy",
    "spain",
    "france",
    "portugal",
    "greece",
    "germany",
    "uk",
    "united kingdom",
    "ireland",
    "usa",
    "united states",
    "mexico",
    "thailand",
    "india",
    "uae",
    "dubai",
    "turkey",
    "croatia",
    "morocco",
    "indonesia",
    "bali",
)

STOP = {
    "the",
    "a",
    "an",
    "and",
    "or",
    "for",
    "client",
    "needs",
    "need",
    "please",
    "booking",
    "book",
    "in",
    "on",
    "from",
    "to",
    "of",
    "with",
    "email",
    "check",
    "internal",
    "availability",
    "available",
    "pricing",
    "price",
    "send",
    "person",
    "nights",
    "night",
    "type",
    "accommodation",
}

TRIGGERS = (
    "check availability",
    "check internal",
    "search inventory",
    "look up booking",
    "lookup booking",
    "what's available",
    "whats available",
    "internal data",
    "internal catalog",
    "check data",
    "search catalog",
)

SHEET_URL_RE = re.compile(
    r"https?://docs\.google\.com/spreadsheets/[^\s)>\"]+",
    re.I,
)
ISO_RE = re.compile(r"\b(\d{4})-(\d{2})-(\d{2})\b")
MONTHS = {
    "jan": 1,
    "january": 1,
    "feb": 2,
    "february": 2,
    "mar": 3,
    "march": 3,
    "apr": 4,
    "april": 4,
    "may": 5,
    "jun": 6,
    "june": 6,
    "jul": 7,
    "july": 7,
    "aug": 8,
    "august": 8,
    "sep": 9,
    "sept": 9,
    "september": 9,
    "oct": 10,
    "october": 10,
    "nov": 11,
    "november": 11,
    "dec": 12,
    "december": 12,
}
RANGE_TEXT_RE = re.compile(
    r"\b(\d{1,2})\s*[–\-to]+\s*(\d{1,2})\s+(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t|tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\b(?:\s+(\d{4}))?",
    re.I,
)
IN_PLACE_RE = re.compile(r"\bin\s+([A-Za-z][A-Za-z ]{1,40}?)(?=\s+(?:from|on|for|the|\d)|,|$)", re.I)

COUNTRY_KEYS = ("country", "location", "destination", "place", "region", "city")
TYPE_KEYS = ("type", "accommodation", "stay", "category", "room_type", "property")
FROM_KEYS = ("check_in", "checkin", "start", "from", "date_from", "available_from", "season_start")
TO_KEYS = ("check_out", "checkout", "end", "to", "date_to", "available_to", "season_end")
NAME_KEYS = ("name", "property", "hotel", "title", "listing")
AVAIL_KEYS = ("available", "availability", "stock", "qty", "rooms", "units")
PRICE_KEYS = ("price", "rate", "cost", "amount", "nightly", "eur", "usd")


def _workspace_root() -> str:
    return os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))


def _parse_date(raw: str) -> Optional[datetime]:
    text = (raw or "").strip()
    if not text:
        return None
    m = ISO_RE.search(text)
    if m:
        try:
            return datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
        except ValueError:
            return None
    m = re.search(
        r"\b(\d{1,2})\s+(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t|tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s+(\d{4})\b",
        text,
        re.I,
    )
    if m:
        month = MONTHS.get(m.group(2).lower()[:3], 0) or MONTHS.get(m.group(2).lower(), 0)
        try:
            return datetime(int(m.group(3)), month, int(m.group(1)))
        except ValueError:
            return None
    return None


def parse_requirement(text: str) -> Dict[str, Any]:
    raw = (text or "").strip()
    low = raw.lower()
    stay = ""
    for kind in STAY_TYPES:
        if re.search(rf"\b{re.escape(kind)}\b", low):
            stay = "bnb" if kind in ("bnb", "b&b") else kind
            break
    place = ""
    for hint in COUNTRY_HINTS:
        if hint in low:
            place = hint
            break
    if not place:
        m = IN_PLACE_RE.search(raw)
        if m:
            place = m.group(1).strip().rstrip(".,")
    start: Optional[datetime] = None
    end: Optional[datetime] = None
    rm = RANGE_TEXT_RE.search(raw)
    if rm:
        month = MONTHS.get(rm.group(3).lower()[:3], 0) or MONTHS.get(rm.group(3).lower(), 0)
        year = int(rm.group(4)) if rm.group(4) else datetime.now().year
        try:
            start = datetime(year, month, int(rm.group(1)))
            end = datetime(year, month, int(rm.group(2)))
        except ValueError:
            start = end = None
    isos = [datetime(int(a), int(b), int(c)) for a, b, c in ISO_RE.findall(raw)]
    if len(isos) >= 2:
        start, end = isos[0], isos[1]
    elif len(isos) == 1 and not start:
        start = isos[0]
        end = start + timedelta(days=2)
    if start and end and end < start:
        start, end = end, start
    tokens = [
        t
        for t in re.findall(r"[a-z0-9]+", low)
        if t not in STOP and len(t) > 1 and t not in STAY_TYPES
    ]
    return {
        "place": place,
        "stay_type": stay,
        "start": start,
        "end": end,
        "tokens": tokens[:24],
        "raw": raw,
    }


def looks_like_internal_query(text: str) -> bool:
    low = (text or "").lower()
    if any(t in low for t in TRIGGERS):
        return True
    if SHEET_URL_RE.search(text or ""):
        return True
    has_date = bool(ISO_RE.search(text or "") or RANGE_TEXT_RE.search(text or ""))
    has_stay = any(re.search(rf"\b{re.escape(k)}\b", low) for k in STAY_TYPES)
    has_place = bool(IN_PLACE_RE.search(text or "") or any(h in low for h in COUNTRY_HINTS))
    return has_date and has_stay and has_place and len(text or "") > 24


def _field(row: Dict[str, str], keys: Iterable[str]) -> str:
    lower = {k.lower().strip(): v for k, v in row.items()}
    for key in keys:
        if key in lower and str(lower[key]).strip():
            return str(lower[key]).strip()
        for actual, val in lower.items():
            if key in actual.replace(" ", "_") and str(val).strip():
                return str(val).strip()
    return ""


def _row_from_mapping(fields: Dict[str, str], source: str) -> Dict[str, Any]:
    blob = " ".join(str(v) for v in fields.values() if v)
    avail_raw = _field(fields, AVAIL_KEYS).lower()
    available = True
    if avail_raw in ("0", "no", "n", "full", "sold", "none", "unavailable"):
        available = False
    elif avail_raw.replace(".", "", 1).isdigit() and float(avail_raw) <= 0:
        available = False
    return {
        "source": source,
        "name": _field(fields, NAME_KEYS) or source,
        "country": _field(fields, COUNTRY_KEYS),
        "stay_type": _field(fields, TYPE_KEYS),
        "date_from": _parse_date(_field(fields, FROM_KEYS)),
        "date_to": _parse_date(_field(fields, TO_KEYS)),
        "available": available,
        "availability_raw": _field(fields, AVAIL_KEYS),
        "price": _field(fields, PRICE_KEYS),
        "blob": blob.lower(),
        "fields": fields,
    }


def _load_csv(path: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    with open(path, "r", encoding="utf-8-sig", errors="replace", newline="") as f:
        sample = f.read(4096)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",\t;")
        except csv.Error:
            dialect = csv.excel
        reader = csv.DictReader(f, dialect=dialect)
        for i, rec in enumerate(reader):
            if i >= MAX_ROWS:
                break
            if not isinstance(rec, dict):
                continue
            clean = {str(k or "").strip(): str(v or "").strip() for k, v in rec.items() if k}
            if any(clean.values()):
                rows.append(_row_from_mapping(clean, os.path.basename(path)))
    return rows


def _load_json(path: str) -> List[Dict[str, Any]]:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        data = json.load(f)
    records: List[Any]
    if isinstance(data, list):
        records = data
    elif isinstance(data, dict):
        records = data.get("records") or data.get("items") or data.get("rows") or [data]
    else:
        return []
    rows = []
    for rec in records[:MAX_ROWS]:
        if isinstance(rec, dict):
            clean = {str(k): str(v) for k, v in rec.items()}
            rows.append(_row_from_mapping(clean, os.path.basename(path)))
    return rows


def _load_xlsx(path: str) -> List[Dict[str, Any]]:
    try:
        import openpyxl
    except ImportError:
        return []
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb.active
        rows_iter = ws.iter_rows(values_only=True)
        header_row = next(rows_iter, None)
        if not header_row:
            return []
        headers = [str(h or f"col{i}").strip() for i, h in enumerate(header_row)]
        out: List[Dict[str, Any]] = []
        for i, vals in enumerate(rows_iter):
            if i >= MAX_ROWS:
                break
            fields = {headers[j]: str(vals[j] or "").strip() for j in range(min(len(headers), len(vals)))}
            if any(fields.values()):
                out.append(_row_from_mapping(fields, os.path.basename(path)))
        return out
    finally:
        wb.close()


def _load_pdf(path: str) -> List[Dict[str, Any]]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return []
    reader = PdfReader(path)
    text = "\n".join((page.extract_text() or "") for page in reader.pages[:20])
    return _rows_from_plain(text, os.path.basename(path))


def _load_docx(path: str) -> List[Dict[str, Any]]:
    try:
        import docx
    except ImportError:
        return []
    document = docx.Document(path)
    table_rows: List[Dict[str, Any]] = []
    for table in document.tables:
        if not table.rows:
            continue
        headers = [c.text.strip() or f"col{i}" for i, c in enumerate(table.rows[0].cells)]
        for row in table.rows[1:]:
            fields = {headers[i]: row.cells[i].text.strip() for i in range(min(len(headers), len(row.cells)))}
            if any(fields.values()):
                table_rows.append(_row_from_mapping(fields, os.path.basename(path)))
    if table_rows:
        return table_rows
    text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    return _rows_from_plain(text, os.path.basename(path))


def _rows_from_plain(text: str, source: str) -> List[Dict[str, Any]]:
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    if not lines:
        return []
    # Try tab/comma tables
    if any("\t" in ln or ln.count(",") >= 2 for ln in lines[:8]):
        sample = "\n".join(lines[:40])
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",\t;")
            reader = csv.DictReader(StringIO("\n".join(lines)), dialect=dialect)
            out = []
            for rec in list(reader)[:MAX_ROWS]:
                clean = {str(k or "").strip(): str(v or "").strip() for k, v in rec.items() if k}
                if any(clean.values()):
                    out.append(_row_from_mapping(clean, source))
            if out:
                return out
        except csv.Error:
            pass
    blob = " ".join(lines)
    return [_row_from_mapping({"text": blob[:4000], "name": source}, source)]


def parse_clipboard_table(text: str, source: str = "google_sheet") -> List[Dict[str, Any]]:
    return _rows_from_plain(text or "", source)


def _catalog_settings(root: Optional[str] = None) -> Dict[str, Any]:
    base = root if root is not None else _workspace_root()
    rb = load_merged_config(base).get("runtime_boundaries") or {}
    path = str(rb.get("internal_data_path") or DEFAULT_REL).strip()
    extra = rb.get("internal_data_extra_roots") or []
    if isinstance(extra, str):
        extra = [extra]
    sheets = rb.get("internal_sheets") or []
    if not isinstance(sheets, list):
        sheets = []
    return {"root": base, "path": path, "extra": list(extra), "sheets": sheets}


def resolve_catalog_dirs(root: Optional[str] = None) -> List[str]:
    settings = _catalog_settings(root)
    base = settings["root"]
    dirs: List[str] = []
    primary = settings["path"]
    abs_primary = primary if os.path.isabs(primary) else os.path.abspath(os.path.join(base, primary))
    if os.path.isdir(abs_primary):
        dirs.append(abs_primary)
    for extra in settings["extra"]:
        extra_s = str(extra or "").strip()
        if not extra_s or extra_s.lower().startswith("http"):
            continue
        abs_extra = extra_s if os.path.isabs(extra_s) else os.path.abspath(os.path.join(base, extra_s))
        if os.path.isdir(abs_extra) and abs_extra not in dirs:
            dirs.append(abs_extra)
    return dirs


def _iter_files(dirs: List[str]) -> List[str]:
    found: List[str] = []
    skip_names = {".git", "__pycache__"}
    for folder in dirs:
        for dirpath, dirnames, filenames in os.walk(folder):
            dirnames[:] = [d for d in dirnames if d not in skip_names and not d.startswith(".")]
            for name in filenames:
                ext = os.path.splitext(name)[1].lower()
                if ext not in SUPPORTED:
                    continue
                path = os.path.join(dirpath, name)
                try:
                    if os.path.getsize(path) > MAX_FILE_BYTES:
                        continue
                except OSError:
                    continue
                found.append(path)
                if len(found) >= MAX_FILES:
                    return found
    return found


def load_catalog_rows(root: Optional[str] = None) -> Tuple[List[Dict[str, Any]], List[str]]:
    dirs = resolve_catalog_dirs(root)
    files = _iter_files(dirs)
    rows: List[Dict[str, Any]] = []
    loaded: List[str] = []
    for path in files:
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext in {".csv", ".tsv"}:
                chunk = _load_csv(path)
            elif ext == ".json":
                chunk = _load_json(path)
            elif ext == ".xlsx":
                chunk = _load_xlsx(path)
            elif ext == ".pdf":
                chunk = _load_pdf(path)
            elif ext == ".docx":
                chunk = _load_docx(path)
            else:
                with open(path, "r", encoding="utf-8", errors="replace") as f:
                    chunk = _rows_from_plain(f.read(MAX_FILE_BYTES), os.path.basename(path))
        except Exception:
            continue
        if chunk:
            loaded.append(os.path.basename(path))
            rows.extend(chunk)
        if len(rows) >= MAX_ROWS:
            break
    return rows[:MAX_ROWS], loaded


def _dates_overlap(req: Dict[str, Any], row: Dict[str, Any]) -> Optional[bool]:
    rs, re_ = req.get("start"), req.get("end")
    ds, de = row.get("date_from"), row.get("date_to")
    if not rs:
        return None
    if not ds and not de:
        return None
    if ds and not de:
        de = ds
    if de and not ds:
        ds = de
    end = re_ or (rs + timedelta(days=1) if rs else None)
    if not rs or not end or not ds or not de:
        return None
    return rs <= de and end >= ds


def score_row(req: Dict[str, Any], row: Dict[str, Any]) -> int:
    score = 0
    blob = str(row.get("blob") or "")
    place = str(req.get("place") or "").lower()
    if place:
        country = str(row.get("country") or "").lower()
        if place in country or place in blob:
            score += 6
        elif country and (country in place or place.split()[0] in country):
            score += 4
        else:
            score -= 2
    stay = str(req.get("stay_type") or "").lower()
    if stay:
        row_stay = str(row.get("stay_type") or "").lower()
        if stay in row_stay or stay in blob:
            score += 5
        else:
            score -= 1
    overlap = _dates_overlap(req, row)
    if overlap is True:
        score += 4
    elif overlap is False:
        score -= 3
    for tok in req.get("tokens") or []:
        if tok in blob:
            score += 1
    if not row.get("available", True):
        score -= 2
    return score


def search_catalog(query: str, *, root: Optional[str] = None, extra_rows: Optional[List[Dict[str, Any]]] = None) -> Dict[str, Any]:
    req = parse_requirement(query)
    rows, files = load_catalog_rows(root)
    if extra_rows:
        rows = list(rows) + list(extra_rows)
    ranked = []
    for row in rows:
        sc = score_row(req, row)
        if sc <= 0 and (req.get("place") or req.get("stay_type")):
            continue
        ranked.append((sc, row))
    ranked.sort(key=lambda x: x[0], reverse=True)
    hits = ranked[:MAX_HITS]
    return {
        "requirement": req,
        "files": files,
        "hits": hits,
        "row_count": len(rows),
    }


def format_catalog_answer(query: str, *, root: Optional[str] = None, extra_rows: Optional[List[Dict[str, Any]]] = None) -> str:
    result = search_catalog(query, root=root, extra_rows=extra_rows)
    req = result["requirement"]
    start = req.get("start")
    end = req.get("end")
    date_bit = ""
    if start:
        date_bit = start.strftime("%Y-%m-%d")
        if end:
            date_bit += " to " + end.strftime("%Y-%m-%d")
    header = [
        "Internal catalog (local files; company data stays on this PC).",
        f"Query: place={req.get('place') or '(any)'} type={req.get('stay_type') or '(any)'} dates={date_bit or '(any)'}",
        f"Crawled {len(result['files'])} file(s), {result['row_count']} rows.",
    ]
    if result["files"]:
        header.append("Files: " + ", ".join(result["files"][:12]))
    else:
        header.append(
            "No catalog files yet. Drop CSV/Excel/PDF/Word into storage/internal_data "
            "(or storage/internal_data/local/) or set runtime_boundaries.internal_data_path."
        )
    available = [(s, r) for s, r in result["hits"] if r.get("available", True)]
    full = [(s, r) for s, r in result["hits"] if not r.get("available", True)]
    lines = header + [""]
    if available:
        lines.append("Available:")
        for i, (sc, row) in enumerate(available, 1):
            lines.append(_format_hit(i, row, sc))
    if full:
        lines.append("Unavailable / full (flagged):")
        for i, (sc, row) in enumerate(full, 1):
            lines.append(_format_hit(i, row, sc))
    if not available and not full:
        lines.append("No matching inventory for that request. Try another country, dates, or stay type.")
    return "\n".join(lines)


def _format_hit(i: int, row: Dict[str, Any], score: int) -> str:
    df = row.get("date_from")
    dt = row.get("date_to")
    dates = ""
    if df:
        dates = df.strftime("%Y-%m-%d")
        if dt:
            dates += " – " + dt.strftime("%Y-%m-%d")
    price = row.get("price") or ""
    avail = row.get("availability_raw") or ("yes" if row.get("available") else "no")
    return (
        f"{i}. {row.get('name') or '(listing)'} — {row.get('country') or '?'} — "
        f"{row.get('stay_type') or '?'} — {dates or 'dates n/a'} — "
        f"stock {avail}"
        + (f" — {price}" if price else "")
        + f" [{row.get('source')}]"
    )


def extract_sheet_url(text: str, *, root: Optional[str] = None) -> str:
    m = SHEET_URL_RE.search(text or "")
    if m:
        return m.group(0).rstrip(".,);")
    settings = _catalog_settings(root)
    low = (text or "").lower()
    for item in settings["sheets"]:
        if not isinstance(item, dict):
            continue
        name = str(item.get("name") or "").strip()
        url = str(item.get("url") or "").strip()
        if url and name and name.lower() in low:
            return url
    if re.search(r"\b(google\s*sheet|gsheet|spreadsheet)\b", low) and settings["sheets"]:
        first = settings["sheets"][0]
        if isinstance(first, dict):
            return str(first.get("url") or "").strip()
    return ""


def sheet_requested(text: str, *, root: Optional[str] = None) -> bool:
    """True only when a Sheet URL is in chat or a configured/named sheet can be resolved."""
    return bool(extract_sheet_url(text, root=root))


def internal_query_plan(query: str, *, root: Optional[str] = None) -> Optional[Dict[str, Any]]:
    """Classify a chat line: local instant answer, or Approve-gated signed-in Sheet copy."""
    if not looks_like_internal_query(query):
        return None
    if sheet_requested(query, root=root):
        url = extract_sheet_url(query, root=root)
        preview = format_catalog_answer(query, root=root)
        payload = json.dumps({"url": url, "query": query})
        return {"kind": "sheet", "url": url, "preview": preview, "payload": payload}
    return {"kind": "instant", "text": format_catalog_answer(query, root=root)}


def capture_signed_in_sheet(url: str, *, run_controller=None, sleep=None) -> Dict[str, Any]:
    """Open a Google Sheet URL in signed-in Chrome/Edge and copy the grid. Desktop — call after Approve."""
    import time as time_mod

    pause = sleep or time_mod.sleep
    try:
        import pyautogui
        import pyperclip
    except ImportError as exc:
        return {"ok": False, "error": f"Desktop capture unavailable: {exc}"}
    from src.execution.landmark_target import focus_window_by_title

    url_s = (url or "").strip()
    if not url_s.startswith("https://docs.google.com/spreadsheets/"):
        return {"ok": False, "error": "Only Google Sheets https URLs are allowed."}

    focused = None
    last = {}
    for title in ("Google Chrome", "Chrome", "Microsoft Edge", "Edge"):
        last = focus_window_by_title(title)
        if last.get("ok"):
            focused = title
            break
    if not focused:
        if run_controller is not None:
            run_controller.enter_stuck(
                last.get("error") or "Chrome or Edge not found.",
                "Open Google Chrome or Edge signed into Google, then Continue.",
            )
            decision = run_controller.wait_while_paused()
            if decision == "stop":
                return {"ok": False, "error": "Stopped while waiting for the browser.", "stop": True}
            if decision != "skip":
                for title in ("Google Chrome", "Chrome", "Microsoft Edge", "Edge"):
                    last = focus_window_by_title(title)
                    if last.get("ok"):
                        focused = title
                        break
        if not focused:
            return {"ok": False, "error": "Open Chrome or Edge (signed into Google) and try again."}

    try:
        pyperclip.copy(url_s)
        pause(0.2)
        pyautogui.hotkey("ctrl", "l")
        pause(0.2)
        pyautogui.hotkey("ctrl", "v")
        pause(0.1)
        pyautogui.press("enter")
        pause(2.8)
        pyautogui.hotkey("ctrl", "a")
        pause(0.25)
        pyautogui.hotkey("ctrl", "c")
        pause(0.35)
        text = pyperclip.paste() or ""
    except Exception as exc:
        return {"ok": False, "error": str(exc)}
    if not text.strip() or text.strip() == url_s:
        return {"ok": False, "error": "Clipboard did not contain a sheet grid. Click inside the sheet, then retry."}
    return {"ok": True, "text": text, "window": focused}


def answer_after_sheet(query: str, clipboard_text: str, *, root: Optional[str] = None) -> str:
    extra = parse_clipboard_table(clipboard_text, source="google_sheet")
    body = format_catalog_answer(query, root=root, extra_rows=extra)
    return body + "\n\n(Includes rows copied from the signed-in Google Sheet after Approve.)"


def redact_catalog_line(text: str) -> str:
    return redact_secrets(text or "")
